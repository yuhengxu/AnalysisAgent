"""从 yuebao 样例目录加载预测表与月报参考内容。"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

import openpyxl
from docx import Document

from app.templates.prediction_table import (
    IMPACT_OPTIONS,
    PREDICTION_FACTORS,
    all_factor_defs,
)
from app.templates.sample_contracts import PREDICTION_SAMPLE_DIR, REPORT_SAMPLE_DIR

_FACTOR_ID_RE = re.compile(r"^\s*(\d+\.\d+)\s+(.+)$")
_PRICE_RANGE_RE = re.compile(
    r"区间.{0,40}?(\d+(?:\.\d+)?)\s*[-–~至]\s*(\d+(?:\.\d+)?)"
)
_PRICE_AVG_RE = re.compile(r"均价[^：:\d]*[：:]?\s*(\d+(?:\.\d+)?)")


def prev_period(year: int, month: int) -> tuple[int, int]:
    if month > 1:
        return year, month - 1
    return year - 1, 12


def prediction_sample_path(year: int, month: int) -> Path:
    return PREDICTION_SAMPLE_DIR / f"{year}年{month}月油价预测分析表.xlsx"


def report_sample_path(year: int, month: int) -> Path | None:
    exact = REPORT_SAMPLE_DIR / f"国际油价月报{year}年第{month}期（总{51 + month}期）.docx"
    if exact.exists():
        return exact
    matches = sorted(REPORT_SAMPLE_DIR.glob(f"国际油价月报{year}年第{month}期*.docx"))
    return matches[0] if matches else None


def find_prediction_sample(year: int, month: int) -> Path | None:
    """优先取上一期样例，若无则向前回溯。"""
    y, m = year, month
    for _ in range(12):
        y, m = prev_period(y, m)
        path = prediction_sample_path(y, m)
        if path.exists():
            return path
    return None


def find_report_sample(outlook_year: int, outlook_month: int) -> Path | None:
    y, m = outlook_year, outlook_month
    for _ in range(12):
        y, m = prev_period(y, m)
        path = report_sample_path(y, m)
        if path and path.exists():
            return path
    return None


def _detect_importance(row: tuple[Any, ...], label_col: int) -> int:
    for cell in row[label_col + 1 : label_col + 6]:
        if cell in (1, 2, 3, 4, 5):
            return int(cell)
    for cell in row[label_col + 1 : label_col + 6]:
        if isinstance(cell, (int, float)) and 1 <= int(cell) <= 5:
            return int(cell)
    return 1


def _detect_impact(row: tuple[Any, ...], label_col: int) -> str:
    for cell in row[label_col + 6 :]:
        text = str(cell).strip() if cell is not None else ""
        if text in IMPACT_OPTIONS:
            return text
    return "持平"


def _parse_price_block(text: str) -> dict[str, Any]:
    lo = hi = avg = None
    m = _PRICE_RANGE_RE.search(text)
    if m:
        lo, hi = float(m.group(1)), float(m.group(2))
    m = _PRICE_AVG_RE.search(text)
    if m:
        avg = float(m.group(1))
    return {"range_low": lo, "range_high": hi, "avg": avg}


def parse_prediction_xlsx(path: Path) -> dict[str, Any]:
    """将 yuebao/prediction 样例 xlsx 解析为结构化 JSON。"""
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = list(ws.iter_rows(values_only=True))
    wb.close()

    label_col = 1
    for row in rows[:8]:
        if not row:
            continue
        for idx, cell in enumerate(row[:6]):
            if cell and re.match(r"^\s*\d+\.\d+", str(cell).strip()):
                label_col = idx
                break

    defs = {d["id"]: d for d in all_factor_defs()}
    by_id: dict[str, dict[str, Any]] = {}
    current_cat = ""
    current_cat_title = ""
    price_blocks: list[dict[str, Any]] = []

    for row in rows:
        if not row or label_col >= len(row):
            continue
        label = str(row[label_col]).strip() if row[label_col] else ""
        if not label:
            continue

        if re.match(r"^\d+\.\S", label) and not re.match(r"^\d+\.\d+", label):
            current_cat_title = label
            for cat in PREDICTION_FACTORS:
                if cat["title"] == label or label in cat["title"]:
                    current_cat = cat["id"]
                    break
            continue

        m = _FACTOR_ID_RE.match(label)
        if m:
            fid, name = m.group(1), m.group(2).strip()
            d = defs.get(fid, {})
            judgment = ""
            for cell in row[label_col + 6 : label_col + 12]:
                if cell and str(cell).strip() not in IMPACT_OPTIONS:
                    judgment = str(cell).strip()
                    break
            by_id[fid] = {
                "category": d.get("category", current_cat),
                "category_title": d.get("category_title", current_cat_title),
                "id": fid,
                "name": d.get("name", name),
                "importance": _detect_importance(row, label_col),
                "judgment": judgment,
                "impact": _detect_impact(row, label_col),
            }
            continue

        if "布伦特" in label and "预测" in label:
            text = " ".join(str(c) for c in row if c)
            block = _parse_price_block(text)
            block["label"] = label
            price_blocks.append(block)

    factors = []
    for d in all_factor_defs():
        item = by_id.get(d["id"], {})
        factors.append(
            {
                "category": d["category"],
                "category_title": d["category_title"],
                "id": d["id"],
                "name": d["name"],
                "importance": item.get("importance", 1),
                "judgment": item.get("judgment", ""),
                "impact": item.get("impact", "持平"),
            }
        )

    month_match = re.search(r"(\d{4})年(\d{1,2})月", path.name)
    sy = int(month_match.group(1)) if month_match else 0
    sm = int(month_match.group(2)) if month_match else 0
    ny, nm = (sy, sm + 1) if sm < 12 else (sy + 1, 1)

    return {
        "sample_file": path.name,
        "sample_year": sy,
        "sample_month": sm,
        "factors": factors,
        "price_forecast": {
            "current_month": {
                "label": price_blocks[0].get("label", f"{sy}年{sm}月份布伦特首行合约价格预测")
                if price_blocks
                else f"{sy}年{sm}月份布伦特首行合约价格预测",
                "range_low": (price_blocks[0] if price_blocks else {}).get("range_low"),
                "range_high": (price_blocks[0] if price_blocks else {}).get("range_high"),
                "avg": (price_blocks[0] if price_blocks else {}).get("avg"),
            },
            "next_month": {
                "label": price_blocks[1].get("label", f"{ny}年{nm}月份布伦特首行合约价格预测")
                if len(price_blocks) > 1
                else f"{ny}年{nm}月份布伦特首行合约价格预测",
                "range_low": (price_blocks[1] if len(price_blocks) > 1 else {}).get("range_low"),
                "range_high": (price_blocks[1] if len(price_blocks) > 1 else {}).get("range_high"),
                "avg": (price_blocks[1] if len(price_blocks) > 1 else {}).get("avg"),
            },
        },
    }


def load_prediction_sample(year: int, month: int) -> dict[str, Any]:
    path = find_prediction_sample(year, month)
    if not path:
        raise FileNotFoundError(f"未找到 {year}年{month}月 预测分析表可参考的 yuebao 样例")
    content = parse_prediction_xlsx(path)
    content["sample_path"] = str(path)
    return content


def parse_report_docx(path: Path) -> dict[str, Any]:
    """从 yuebao/yuebao 月报 docx 提取文本与表格，供大模型仿写。"""
    doc = Document(str(path))
    paragraphs: list[dict[str, str]] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name if p.style else "") or ""
        paragraphs.append({"style": style, "text": text})

    tables: list[dict[str, Any]] = []
    for tbl in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in tbl.rows]
        if rows:
            tables.append({"rows": rows})

    issue_match = re.search(r"(\d{4})年第(\d{1,2})期", path.name)
    return {
        "sample_file": path.name,
        "sample_year": int(issue_match.group(1)) if issue_match else None,
        "sample_month": int(issue_match.group(2)) if issue_match else None,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def load_report_sample(outlook_year: int, outlook_month: int) -> dict[str, Any]:
    path = find_report_sample(outlook_year, outlook_month)
    if not path:
        raise FileNotFoundError(
            f"未找到 {outlook_year}年{outlook_month}月 月报可参考的 yuebao 样例"
        )
    content = parse_report_docx(path)
    content["sample_path"] = str(path)
    return content
