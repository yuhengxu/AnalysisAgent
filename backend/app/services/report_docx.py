"""月报 Word 导出：基于样例 docx 模板，对齐字号、字体、行距、标题与表格格式。"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.templates.sample_contracts import REPORT_CHART_ANCHORS, REPORT_SAMPLE_PATH, REPORT_TABLE_ANCHORS

# 样例月报（保留页边距、内置样式与 Table Grid 表格样式）
SAMPLE_TEMPLATE_PATH = REPORT_SAMPLE_PATH

# 与样例一致的版式常量（单位 pt）
FONT_ASCII = "Times New Roman"
INDENT_BODY = Pt(28)
INDENT_H2 = Pt(28.1)
INDENT_H1 = Pt(32)
SPACE_SECTION = Pt(7.83)
SPACE_SUMMARY_BEFORE = Pt(31.25)
SIZE_COVER_ORG = 22
SIZE_COVER_TITLE = 28
SIZE_BODY = 14
SIZE_TABLE_TITLE = 15
SIZE_TABLE_CELL = 11
SIZE_SOURCE = 10


def build_monthly_report_document(
    content: dict[str, Any],
    charts: dict[str, str] | None = None,
    *,
    review_year: int | None = None,
    review_month: int | None = None,
) -> Document:
    """根据结构化月报 JSON 生成与样例格式一致的 Word 文档。"""
    charts = charts or {}
    doc = _load_template()
    _add_cover(doc, content.get("cover", {}))
    _add_summary(doc, str(content.get("summary", "")))
    tables = content.get("tables", {})
    chart_anchors: dict[str, list[dict[str, Any]]] = {}
    for item in REPORT_CHART_ANCHORS:
        chart_anchors.setdefault(item["section_id"], []).append(item)
    for sec in content.get("sections", []):
        level = sec.get("level", 2)
        title = sec.get("title", "")
        if level == 1:
            _add_heading1(doc, title)
        else:
            _add_heading2(doc, title)
            if sec.get("content"):
                _add_body_paragraphs(doc, sec["content"])
        sec_id = sec.get("id")
        for chart in chart_anchors.get(sec_id, []):
            _add_chart(
                doc,
                chart,
                charts.get(chart["id"]),
                content=content,
                review_year=review_year,
                review_month=review_month,
            )
        for table_key, anchor_id in REPORT_TABLE_ANCHORS.items():
            if sec_id == anchor_id:
                _add_table(doc, tables.get(table_key))
    _add_approval(doc, content.get("approval", {}))
    _add_table(doc, tables.get("table_distribution"))
    return doc


def _load_template() -> Document:
    if not SAMPLE_TEMPLATE_PATH.exists():
        return Document()
    doc = Document(str(SAMPLE_TEMPLATE_PATH))
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            body.remove(child)
    return doc


def _set_run_font(
    run,
    east_asia: str,
    *,
    size_pt: float | None = None,
    bold: bool | None = None,
) -> None:
    run.font.name = FONT_ASCII
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), FONT_ASCII)
    r_fonts.set(qn("w:hAnsi"), FONT_ASCII)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), east_asia if east_asia in ("黑体", "仿宋") else FONT_ASCII)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _add_cover(doc: Document, cover: dict[str, Any]) -> None:
    org = cover.get("org", "")
    if org:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(org)
        _set_run_font(run, "方正小标宋简体", size_pt=SIZE_COVER_ORG, bold=True)

    title = cover.get("title", "")
    if title:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        _set_run_font(run, "华文中宋", size_pt=SIZE_COVER_TITLE, bold=True)

    issue = cover.get("issue", "")
    if issue:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = Pt(SIZE_COVER_TITLE)
        run = para.add_run(issue)
        _set_run_font(run, "仿宋")

    dept = cover.get("dept", "")
    date = cover.get("date", "")
    if dept or date:
        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = Pt(SIZE_COVER_TITLE)
        run = para.add_run(f"{dept}            {date}".strip())
        _set_run_font(run, "仿宋", size_pt=SIZE_BODY)


def _add_summary(doc: Document, summary: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = SPACE_SUMMARY_BEFORE
    para.paragraph_format.space_after = SPACE_SECTION
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = para.add_run("内容摘要")
    _set_run_font(run, "黑体", bold=True)

    for block in _split_paragraphs(summary):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = INDENT_BODY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        r = p.add_run(block)
        _set_run_font(r, "仿宋", size_pt=SIZE_BODY)


def _add_heading1(doc: Document, title: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = INDENT_H1
    para.paragraph_format.space_after = SPACE_SECTION
    run = para.add_run(title)
    _set_run_font(run, "黑体", bold=True)


def _add_heading2(doc: Document, title: str) -> None:
    para = doc.add_paragraph(style="Heading 2")
    pf = para.paragraph_format
    pf.first_line_indent = INDENT_H2
    pf.space_before = SPACE_SECTION
    pf.space_after = SPACE_SECTION
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    run = para.add_run(title)
    _set_run_font(run, "仿宋", size_pt=SIZE_BODY)


def _add_body_paragraphs(doc: Document, text: str) -> None:
    for block in _split_paragraphs(text):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = INDENT_BODY
        run = para.add_run(block)
        _set_run_font(run, "方正仿宋简体", size_pt=SIZE_BODY)


def _add_table(doc: Document, tbl: dict[str, Any] | None) -> None:
    if not tbl or not tbl.get("headers"):
        return

    title = tbl.get("title", "")
    if title:
        cap = _add_caption_paragraph(doc)
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap.paragraph_format.line_spacing = Pt(SIZE_COVER_TITLE)
        run = cap.add_run(title)
        _set_run_font(run, "黑体", size_pt=SIZE_TABLE_TITLE)

    headers = tbl["headers"]
    rows = tbl.get("rows", [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, str(header), center=True)

    for row in rows:
        cells = table.add_row().cells
        for i in range(len(headers)):
            value = str(row[i]) if i < len(row) else ""
            _set_cell_text(cells[i], value, center=True)

    source = tbl.get("source", "")
    if source:
        src = _add_caption_paragraph(doc)
        src.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        src.paragraph_format.line_spacing = Pt(SIZE_COVER_TITLE)
        run = src.add_run(f"数据来源：{source}")
        _set_run_font(run, "黑体", size_pt=SIZE_SOURCE)


def _add_chart(
    doc: Document,
    chart: dict[str, Any],
    path: str | None,
    *,
    content: dict[str, Any],
    review_year: int | None = None,
    review_month: int | None = None,
) -> None:
    from app.services.chart_export import ChartExportService

    if not ChartExportService.is_valid_chart_file(path):
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=Inches(6.4))
    title = ChartExportService.format_chart_title(
        chart,
        review_year=review_year,
        review_month=review_month,
        content=content,
    )
    if title:
        cap = _add_caption_paragraph(doc)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(title)
        _set_run_font(run, "黑体", size_pt=SIZE_TABLE_TITLE)
    source = chart.get("source", "")
    if source:
        src = _add_caption_paragraph(doc)
        src.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = src.add_run(f"数据来源：{source}")
        _set_run_font(run, "黑体", size_pt=SIZE_SOURCE)


def _add_caption_paragraph(doc: Document):
    try:
        return doc.add_paragraph(style="图表标题")
    except KeyError:
        return doc.add_paragraph()


def _set_cell_text(cell, text: str, *, center: bool = False) -> None:
    cell.text = text
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        _set_run_font(run, "方正仿宋简体", size_pt=SIZE_TABLE_CELL)


def _add_approval(doc: Document, approval: dict[str, Any]) -> None:
    doc.add_paragraph()
    for key in ("author", "reviewer", "approver", "signer"):
        text = approval.get(key, "")
        if not text:
            continue
        para = doc.add_paragraph()
        if key != "signer":
            para.paragraph_format.first_line_indent = INDENT_H1
        run = para.add_run(text)
        _set_run_font(run, "方正小标宋简体", size_pt=SIZE_BODY)


def _split_paragraphs(text: str) -> list[str]:
    blocks: list[str] = []
    for line in text.replace("\r\n", "\n").split("\n"):
        stripped = line.strip()
        if stripped:
            blocks.append(stripped)
    return blocks
